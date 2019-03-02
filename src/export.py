#! python
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt

class DocxWriter:
    # Full file name.
    __fileName  = None
    # Document instance.
    __document  = None
    # Current paragraph
    __paragraph = None

    # Create a new docx file, the first step.
    def create(self, fullName):
        self.__reset()
        self.__fileName = "%s.docx" % (fullName)
        self.__document = Document()

    # Save edition, the last step.
    def save(self):
        self.__document.save(self.__fileName)
        self.__clear()

    def addTitle(self, content, level = 0):
        self.__document.add_heading(content, level)

    def addParagraph(self, content = None, alignment = 0):
        self.__paragraph = self.__document.add_paragraph(content)
        self.__paragraph.alignment = alignment

    def addText(self, content, bold = False, italic = False):
        run = self.__paragraph.add_run(content)
        run.bold = bold
        run.italic = italic

    def addPicture(self, fullName, width, height):
        run = self.__paragraph.add_run()
        run.add_picture(fullName, width = Pt(0.5 * width), height = Pt(0.5 * height))

    def __clear(self):
        del self.__fileName
        del self.__document
        del self.__paragraph
